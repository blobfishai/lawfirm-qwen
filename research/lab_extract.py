#!/usr/bin/env python3
"""Extract text from a Harvey LAB task's document corpus.

LAB ships its evidence as real office files — 33,954 .docx, 10,575 .xlsx,
5,169 .eml, 1,091 .pptx (research/answers/data/lab-corpus.json). Our world
stores documents as text rows, so hosting a LAB task means reading the real
bytes rather than paraphrasing them.

Usage:
    python3 research/lab_extract.py <task-dir> [--json] [--max-chars N]

Prints each document's extracted text, or a JSON array with --json.
Formats: .docx (python-docx, paragraphs + tables), .xlsx (openpyxl, sheet
grids), .eml (stdlib email, headers + plain body), .txt/.md (verbatim).
.pptx and .pdf are reported as unsupported rather than silently skipped.
"""
from __future__ import annotations

import argparse
import json
import sys
from email import policy
from email.parser import BytesParser
from pathlib import Path


def from_docx(p: Path) -> str:
    import docx
    d = docx.Document(str(p))
    out = [para.text for para in d.paragraphs if para.text.strip()]
    for ti, table in enumerate(d.tables):
        out.append(f"\n[TABLE {ti + 1}]")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


def from_xlsx(p: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(p), data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"\n[SHEET: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            if row is None:
                continue
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                out.append(" | ".join(cells).rstrip(" |"))
    wb.close()
    return "\n".join(out)


def from_eml(p: Path) -> str:
    msg = BytesParser(policy=policy.default).parse(p.open("rb"))
    hdr = [f"From: {msg.get('From', '')}", f"To: {msg.get('To', '')}",
           f"Cc: {msg.get('Cc', '')}", f"Date: {msg.get('Date', '')}",
           f"Subject: {msg.get('Subject', '')}"]
    try:
        body = msg.get_body(preferencelist=("plain", "html"))
        text = body.get_content() if body else ""
    except Exception:
        text = ""
    return "\n".join(h for h in hdr if h.split(": ", 1)[1]) + "\n\n" + text


EXTRACTORS = {".docx": from_docx, ".xlsx": from_xlsx, ".eml": from_eml,
              ".txt": lambda p: p.read_text(errors="replace"),
              ".md": lambda p: p.read_text(errors="replace")}


def extract_dir(docs: Path, max_chars: int | None = None) -> list[dict]:
    rows = []
    for p in sorted(docs.rglob("*")):
        if not p.is_file():
            continue
        fn = EXTRACTORS.get(p.suffix.lower())
        if fn is None:
            rows.append({"file": str(p.relative_to(docs)), "ext": p.suffix.lower(),
                         "ok": False, "error": "unsupported format", "text": ""})
            continue
        try:
            text = fn(p)
        except Exception as e:  # a failed parse is recorded, never silently dropped
            rows.append({"file": str(p.relative_to(docs)), "ext": p.suffix.lower(),
                         "ok": False, "error": f"{type(e).__name__}: {e}"[:200], "text": ""})
            continue
        if max_chars and len(text) > max_chars:
            text = text[:max_chars]
        rows.append({"file": str(p.relative_to(docs)), "ext": p.suffix.lower(),
                     "ok": True, "chars": len(text), "text": text})
    return rows


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("task_dir")
    ap.add_argument("--json", action="store_true")
    ap.add_argument("--max-chars", type=int, default=None)
    a = ap.parse_args()
    docs = Path(a.task_dir) / "documents"
    if not docs.is_dir():
        print(f"no documents/ under {a.task_dir}", file=sys.stderr)
        return 1
    rows = extract_dir(docs, a.max_chars)
    if a.json:
        print(json.dumps(rows, indent=1))
    else:
        for r in rows:
            head = f"===== {r['file']}  ({r['ext']}) ====="
            print(head)
            print(r["text"] if r["ok"] else f"  [UNPARSED: {r['error']}]")
            print()
    bad = [r for r in rows if not r["ok"]]
    print(f"-- {len(rows) - len(bad)}/{len(rows)} parsed"
          + (f"; unparsed: {[r['file'] for r in bad]}" if bad else ""), file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
